#!/usr/bin/env python3
# scripts/kb/_docx2txt.py —— python-docx 桥接：把 .docx 正文抽取为纯文本到 stdout
import sys
from docx import Document

def main():
    if len(sys.argv) < 2:
        sys.stderr.write("usage: _docx2txt.py <file.docx>\n")
        sys.exit(2)
    path = sys.argv[1]
    try:
        doc = Document(path)
    except Exception as e:
        sys.stderr.write(f"docx 打开失败: {e}\n")
        sys.exit(1)
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    # 表格文本也纳入
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    sys.stdout.write("\n\n".join(parts))

if __name__ == "__main__":
    main()
